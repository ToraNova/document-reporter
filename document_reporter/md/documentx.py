# renders a docx from markdown
# using mistletoe for markdown parsing and python-docx for docx rendering

import os
import re
from urllib.parse import urlparse

from itertools import chain
from mistletoe import block_token, span_token
from mistletoe.base_renderer import BaseRenderer

from docx import Document
from docx.shared import Pt, Inches, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from ..util.docx_helper import format_paragraph, format_run, format_table, insert_hrule, insert_hyperlink, parse_sizespec, assign_numbering, make_figure_caption, shade_cell, delete_paragraph
from .format_tag import HorizontalRuleTag, LineBreakTag, PageBreakTag, ColorTag, BoldTag, ItalicTag, UnderlineTag, StrongTag, EmphasisTag, StrikethroughTag, FontTag, ImageTag, CellTag
from .format_tag import CommentBlockTag, AlignBlockTag, TableBlockTag, ParagraphBlockTag

class Renderer(BaseRenderer):
    def __init__(self, docx_template=None, rel_root=None, *extras):
        '''
        renders a .docx using docx_template as the template file
        rel_root determine the path to include relative path resources (e.g., images)
        '''
        self._suppress_ptag_stack = [False]
        super().__init__(*chain(([
            HorizontalRuleTag, LineBreakTag, PageBreakTag,
            ColorTag, BoldTag, ItalicTag, UnderlineTag, StrongTag, EmphasisTag, StrikethroughTag, FontTag, ImageTag, CellTag,
            CommentBlockTag, AlignBlockTag, TableBlockTag, ParagraphBlockTag,
        ]), extras))
        self.rel_root = os.getcwd() if rel_root is None else rel_root
        self.docx_template = docx_template

    def render_inner(self, token):
        # inner rendering not supported on DocxRenderer
        # raise NotImplementedError('unsupported render_inner function called')
        for c in token.children:
            self.render(c)

    def render_list(self, token):
        # https://github.com/miyuchina/mistletoe/blob/94022647cd9d80e242db5c93a6567e3155b468bc/mistletoe/block_token.py#L447
        ordered = token.start is not None
        level = self.list_level
        self.list_level += 1

        # TODO: fix this abstract_numId map with a more robust implementation
        anid_map = {
                0: 7, # numId 5, absnumId 7
                1: 3, # numId 6, absnumId 3
                2: 2, # numId 7, absnumId 2
        }

        if ordered:
            stpl = 'List Number '
        else:
            stpl = 'List Bullet '

        if level > 0:
            stpl += str(level+1)
        style = stpl.strip()

        for idx, c in enumerate(token.children):
            number = token.start + idx if ordered else None
            # if mistletoe changed the way leader/prepend is computed, this must changed as well

            tos = len(self.paras)
            for ic in c.children:
                ic.docx_style = style
                self.render(ic)
            #self.populate_and_format_paras(c, style=style)

            added = len(self.paras) - tos
            #if added > 0 and ordered:
            if added > 0 and ordered and idx == 0:
                assign_numbering(self.docx, self.paras[tos], anid_map[level], number)

        self.list_level -= 1

    def render_list_item(self, token):
        # handled by list
        #self.render_inner(token)
        raise NotImplementedError("list_item is handled by list rendering")

    def render_line_break(self, token):
        self.runs[-1].add_break()

    def render_link(self, token):
        # TODO: allow alt-txt in links, see mistletoe docs on parser issues (token.title is empty!)
        if len(token.title) < 1:
            ut = token.target
        else:
            ut = token.title
        insert_hyperlink(self.paras[-1], ut, token.target)

    def render_auto_link(self, token):
        insert_hyperlink(self.paras[-1], token.target, token.target)

    def render_image(self, token):
        if token.src.startswith('/'):
            img_path = token.src
        else:
            img_path = os.path.join(self.rel_root, token.src)

        self.runs.append(self.paras[-1].add_run())
        self.inline_shapes.append(self.runs[-1].add_picture(img_path))

        # skip caption if no caption string
        if len(token.title.strip()) < 1:
            return

        tcpar = self.docx.add_paragraph(style='Caption').clear()
        format_paragraph(tcpar, align=WD_ALIGN_PARAGRAPH.CENTER)
        #_hdguse = 1 # use hdg lvl 1 count as index
        #if _hdguse not in self.hdg_count:
        #    predot_num = ''
        #else:
        #    predot_num = f'{self.hdg_count[_hdguse]}-'
        #make_figure_caption(tcpar.add_run(f'Figure {predot_num}'))
        # TODO: allow user to choose whether to include heading
        make_figure_caption(tcpar.add_run(f'Figure '), 1)
        tcrun = tcpar.add_run(f': {token.title}')

    def render_image_tag(self, token):
        up = urlparse(token.format_value_raw)

    def render_raw_text(self, token):
        # add run to last added paragraph
        self.runs.append(self.paras[-1].add_run())
        self.runs[-1].add_text(token.content)

    def render_document(self, token):
        # create document from template
        if self.docx_template is None:
            self.docx = Document()
        else:
            self.docx = Document(self.docx_template)
            delete_paragraph(self.docx.paragraphs[-1])

        # paragraph and run stack
        self.list_level = 0 # default root level
        self.runs = []
        self.paras = self.docx.paragraphs.copy()
        self.tables = self.docx.tables.copy()
        self.hdg_count = {} # for heading number in caption use
        self.cells = []
        self.inline_shapes = []
        self.render_inner(token)

        # save document
        #self.docx.save(self.out_path)
        return self.docx

    def populate_and_format_runs(self, token, **kwargs):
        # format_tag allows nested formatting <b><i>test</i></b>, italic and bold
        tos = len(self.runs) # top of stack
        self.render_inner(token)
        added = len(self.runs) - tos # new top of stack
        # apply format to all added elements during inner render
        for i in range(added):
            format_run(self.runs[-(i+1)], **kwargs)

    def populate_and_format_paras(self, token, **kwargs):
        tos = len(self.paras)
        self.render_inner(token)
        added = len(self.paras) - tos
        # apply format to all added elements during inner render
        for i in range(added):
            format_paragraph(self.paras[-(i+1)], **kwargs)

    def populate_and_format_tables(self, token, **kwargs):
        tos = len(self.tables) # top of stack
        self.render_inner(token)
        added = len(self.tables) - tos # new top of stack
        # add style to all added table within this block
        for i in range(added):
            format_table(self.tables[-(i+1)], **kwargs)

    def populate_and_resize_image(self, token, width, height=None):
        tos = len(self.inline_shapes)
        self.render_inner(token)
        added = len(self.inline_shapes) - tos
        for i in range(added):
            target = self.inline_shapes[-(i+1)]
            if height is not None:
                target.height = height
            else:
                # maintain aspect ratio
                sf = float(width) / float(target.width)
                target.height = round(target.height * sf)

            target.width = width

    def render_color_tag(self, token):
        self.populate_and_format_runs(token, color=RGBColor.from_string(token.format_value))

    def render_bold_tag(self, token):
        self.populate_and_format_runs(token, bold=True)

    def render_strong(self, token):
        self.populate_and_format_runs(token, style='Strong')

    def render_strong_tag(self, token):
        self.populate_and_format_runs(token, style='Strong')

    def render_italic_tag(self, token):
        self.populate_and_format_runs(token, italic=True)

    def render_emphasis(self, token):
        self.populate_and_format_runs(token, style='Emphasis')

    def render_emphasis_tag(self, token):
        self.populate_and_format_runs(token, style='Emphasis')

    def render_underline_tag(self, token):
        self.populate_and_format_runs(token, underline=True)

    def render_strikethrough_tag(self, token):
        self.populate_and_format_runs(token, strike=True)

    def render_strikethrough(self, token):
        self.populate_and_format_runs(token, strike=True)

    def render_font_tag(self, token):
        _name = token.format['name'] if 'name' in token.format else None
        _size = parse_sizespec(token.format['size']) if 'size' in token.format else None
        self.populate_and_format_runs(token, size=_size, name=_name)

    def render_heading(self, token):
        # assume that heading has no additional child
        if token.level not in self.hdg_count:
            self.hdg_count[token.level] = 1
        else:
            self.hdg_count[token.level] += 1
        self.paras.append(self.docx.add_heading(level=token.level).clear())
        self.render_inner(token)

    def render_paragraph(self, token):
        # for all span tokens
        if hasattr(token, 'docx_style') and isinstance(token.docx_style, str):
            self.paras.append(self.docx.add_paragraph(style=token.docx_style).clear())
        else:
            self.paras.append(self.docx.add_paragraph().clear())
        self.render_inner(token)

    def render_align_block_tag(self, token):
        map = {
            'center': WD_ALIGN_PARAGRAPH.CENTER,
            'left': WD_ALIGN_PARAGRAPH.LEFT,
            'right': WD_ALIGN_PARAGRAPH.RIGHT,
            'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        if token.format_value not in map:
            raise KeyError(f'no such alignment type: {token.format_value} (center, left, right, justify)')
        self.populate_and_format_paras(token, align=map[token.format_value])

    def render_horizontal_rule_tag(self, token):
        insert_hrule(self.paras[-1], token.format_value)

    def render_thematic_break(self, token):
        insert_hrule(self.paras[-1])

    def render_page_break_tag(self, token):
        self.docx.add_page_break()

    def render_line_break_tag(self, token):
        #self.runs[-1].text += '\n'
        if len(self.runs) < 1:
            self.runs.append(self.paras[-1].add_run())
        self.runs[-1].add_break()

    def render_table(self, token):
        if len(token.children) < 1:
            return
        ncol = len(token.children[0].children)
        self.tables.append(self.docx.add_table(0, ncol))

        if hasattr(token, 'header') and token.header is not None:
            token.header.is_header = True
            self.render(token.header)

        for row in token.children:
            row.is_header = False
            self.render(row)

    def render_table_row(self, token):
        row = self.tables[-1].add_row()
        for cidx, col in enumerate(token.children):
            self.cells.append(row.cells[cidx])
            self.render(col)

    def render_table_cell(self, token):
        self.paras.append(self.cells[-1].paragraphs[0])
        self.render_inner(token)

    def render_image_tag(self, token):
        _width = parse_sizespec(token.format['width']) if 'width' in token.format else None
        self.populate_and_resize_image(token, width=_width)

    def render_cell_tag(self, token):
        if len(self.cells) < 1:
            return
        if 'color' in token.format:
            _color = token.format['color'].strip('#')
            shade_cell(self.cells[-1], _color)
        self.render_inner(token)

    def render_table_block_tag(self, token):
        _style = token.format['style'] if 'style' in token.format else None
        _colwidths = None
        if 'colwidths' in token.format:
            _colwidth_raw = token.format['colwidths'].split(';')
            _colwidths = []
            for cr in _colwidth_raw:
                _colwidths.append(parse_sizespec(cr.strip()))
        self.populate_and_format_tables(token, colwidths=_colwidths, style=_style)

    def render_paragraph_block_tag(self, token):
        _spacing = parse_sizespec(token.format['spacing']) if 'spacing' in token.format else None
        _before = parse_sizespec(token.format['before']) if 'before' in token.format else None
        _after = parse_sizespec(token.format['after']) if 'after' in token.format else None
        _style = token.format['style'] if 'style' in token.format else None

        _align = None
        if 'align' in token.format:
            map = {
                'center': WD_ALIGN_PARAGRAPH.CENTER,
                'left': WD_ALIGN_PARAGRAPH.LEFT,
                'right': WD_ALIGN_PARAGRAPH.RIGHT,
                'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            if token.format['align'] not in map:
                raise KeyError(f'no such alignment type: {token.format_value} (center, left, right, justify)')
            _align = map[token.format['align']]

        self.populate_and_format_paras(token, style=_style, spacing=_spacing, before=_before, after=_after, align=_align)

    def render_comment_block_tag(self, token):
        # do nothing
        pass
